from __future__ import annotations

from pathlib import Path
import fitz
from docx import Document

from src.common.text_utils import normalize_text
from src.models import SourceDocument, TextBlock


def load_documents(inspection_path: str, thermal_path: str) -> list[SourceDocument]:
    inspection = _load_single(inspection_path, "inspection")
    thermal = _load_single(thermal_path, "thermal")
    return [inspection, thermal]


def _load_single(file_path: str, doc_type: str) -> SourceDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text_blocks = _load_pdf_text(path)
    elif suffix == ".docx":
        text_blocks = _load_docx_text(path)
    else:
        raise ValueError(f"Unsupported document type: {suffix}")

    return SourceDocument(
        doc_id=path.stem,
        doc_type=doc_type,
        file_path=str(path),
        text_blocks=text_blocks,
    )


def _load_pdf_text(path: Path) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    doc = fitz.open(str(path))
    try:
        for idx, page in enumerate(doc, start=1):
            text = normalize_text(page.get_text("text"))
            if text:
                blocks.append(TextBlock(page=idx, text=text))
    finally:
        doc.close()
    return blocks


def _load_docx_text(path: Path) -> list[TextBlock]:
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    text = normalize_text(text)
    if not text:
        return []
    return [TextBlock(page=1, text=text)]
